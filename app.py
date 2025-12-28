from docx import Document
from datetime import datetime, timedelta
import imaplib
import email
from email.header import decode_header
import requests
import os
import re
import time
from flask import Flask, render_template, jsonify, request, redirect, session, flash, send_file
import logging
import sqlite3
import json
import traceback
import threading
import bcrypt
from functools import wraps
import io
import csv
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from fpdf import FPDF

# Initialize Flask app FIRST
app = Flask(__name__)

# Generate a secure secret key (store this in environment variable in production)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')

# Default admin credentials (change these in production)
DEFAULT_USERNAME = os.getenv('DASHBOARD_USERNAME', 'admin')
DEFAULT_PASSWORD = os.getenv('DASHBOARD_PASSWORD', 'admin123')

# Hash the password at module level (no global declaration needed in functions)
HASHED_PASSWORD = bcrypt.hashpw(DEFAULT_PASSWORD.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
print(f"üîê Password hash initialized for user: {DEFAULT_USERNAME}")

def get_db_path():
    """Get database path that works for both web service and cron job"""
    # On Render, use /tmp directory which is shared between services
    if os.getenv('RENDER'):
        return '/tmp/email_summaries.db'
    else:
        return 'email_summaries.db'

# ==================== AUTHENTICATION DECORATORS ====================

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session:
            return redirect('/login?next=' + request.path)
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session or session.get('username') != DEFAULT_USERNAME:
            return redirect('/login?next=' + request.path)
        return f(*args, **kwargs)
    return decorated_function

# ==================== AUTHENTICATION ROUTES ====================

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()
        next_page = request.args.get('next', '/dashboard')
        
        # Verify credentials
        if username == DEFAULT_USERNAME:
            # Check password
            try:
                if bcrypt.checkpw(password.encode('utf-8'), HASHED_PASSWORD.encode('utf-8')):
                    session['logged_in'] = True
                    session['username'] = username
                    session['login_time'] = datetime.now().isoformat()
                    
                    # Set session to expire after 24 hours
                    session.permanent = True
                    app.permanent_session_lifetime = timedelta(hours=24)
                    
                    print(f"üîê User '{username}' logged in successfully")
                    return redirect(next_page)
            except Exception as e:
                print(f"üîê Login error: {e}")
        
        # Invalid credentials
        flash('Invalid username or password', 'error')
        return render_template('login.html', error='Invalid credentials')
    
    # GET request - show login form
    return render_template('login.html')

@app.route('/logout')
def logout():
    """Logout user"""
    username = session.get('username', 'Unknown')
    session.clear()
    print(f"üîê User '{username}' logged out")
    flash('You have been logged out successfully', 'success')
    return redirect('/login')

@app.route('/change-password', methods=['GET', 'POST'])
@admin_required
def change_password():
    """Change password page (admin only)"""
    global HASHED_PASSWORD  # Declare as global since we need to modify it
    
    if request.method == 'POST':
        current_password = request.form.get('current_password', '').strip()
        new_password = request.form.get('new_password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        
        # Validate current password
        try:
            if not bcrypt.checkpw(current_password.encode('utf-8'), HASHED_PASSWORD.encode('utf-8')):
                flash('Current password is incorrect', 'error')
                return render_template('change_password.html')
        except Exception as e:
            flash('Error validating current password', 'error')
            return render_template('change_password.html')
        
        # Validate new password
        if new_password != confirm_password:
            flash('New passwords do not match', 'error')
            return render_template('change_password.html')
        
        if len(new_password) < 8:
            flash('New password must be at least 8 characters', 'error')
            return render_template('change_password.html')
        
        # Update password
        try:
            HASHED_PASSWORD = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        except Exception as e:
            flash('Error hashing new password', 'error')
            return render_template('change_password.html')
        
        # Log out all sessions
        session.clear()
        
        print("üîê Password changed successfully")
        flash('Password changed successfully. Please log in again.', 'success')
        return redirect('/login')
    
    return render_template('change_password.html')

# ==================== DATABASE INITIALIZATION ====================

def init_db():
    """Initialize SQLite database - CALL THIS BEFORE ANY DATABASE OPERATIONS"""
    db_path = get_db_path()
    print(f"üìÅ Initializing database at: {db_path}")
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    
    # Drop tables if they exist (for fresh start)
    c.execute('DROP TABLE IF EXISTS summary_runs')
    c.execute('DROP TABLE IF EXISTS email_data')
    
    c.execute('''
        CREATE TABLE IF NOT EXISTS summary_runs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            run_date TEXT NOT NULL,
            total_emails INTEGER,
            processed_emails INTEGER,
            success_rate REAL,
            deepseek_tokens INTEGER,
            status TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    c.execute('''
        CREATE TABLE IF NOT EXISTS email_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            run_id INTEGER,
            email_number INTEGER,
            sender TEXT,
            receiver TEXT,
            subject TEXT,
            summary TEXT,
            email_date TEXT,
            processed_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (run_id) REFERENCES summary_runs (id)
        )
    ''')
    
    # Create indexes for better performance
    c.execute('CREATE INDEX IF NOT EXISTS idx_run_id ON email_data (run_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_email_number ON email_data (email_number)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_email_date ON email_data (email_date)')
    
    conn.commit()
    conn.close()
    print(f"‚úÖ Database initialized at: {db_path}")

# Initialize database immediately
init_db()

# ==================== EXPORT ROUTES ====================

@app.route('/api/export-data', methods=['POST'])
@login_required
def export_data():
    """Export email data in various formats"""
    try:
        data = request.json
        format_type = data.get('format', 'csv')
        filename = data.get('filename', 'email_summaries')
        
        # Get filters from request
        filters = data.get('filters', {})
        
        # Get email data with filters
        email_data = get_filtered_email_data(filters)
        
        if not email_data:
            return jsonify({"error": "No data to export"}), 400
        
        if format_type == 'csv':
            return export_csv(email_data, filename)
        elif format_type == 'json':
            return export_json(email_data, filename)
        elif format_type == 'word':
            return export_word(email_data, filename)
        elif format_type == 'pdf':
            return export_pdf(email_data, filename)
        else:
            return jsonify({"error": "Unsupported format"}), 400
            
    except Exception as e:
        print(f"‚ùå Export error: {e}")
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

def get_filtered_email_data(filters):
    """Get email data with optional filters"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Build query based on filters
        query = '''
            SELECT email_number, sender, receiver, subject, summary, email_date 
            FROM email_data 
            WHERE 1=1
        '''
        params = []
        
        # Apply date filter
        date_range = filters.get('dateRange')
        if date_range == 'today':
            today = datetime.now().strftime('%Y-%m-%d')
            query += " AND DATE(email_date) = ?"
            params.append(today)
        elif date_range == 'week':
            week_ago = (datetime.now() - timedelta(days=7)).strftime('%Y-%m-%d')
            query += " AND DATE(email_date) >= ?"
            params.append(week_ago)
        elif date_range == 'month':
            month_ago = (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d')
            query += " AND DATE(email_date) >= ?"
            params.append(month_ago)
        elif date_range == 'custom' and filters.get('startDate') and filters.get('endDate'):
            query += " AND DATE(email_date) BETWEEN ? AND ?"
            params.extend([filters['startDate'], filters['endDate']])
        
        # Apply sender/receiver filters
        sender = filters.get('sender')
        if sender:
            query += " AND sender LIKE ?"
            params.append(f'%{sender}%')
        
        receiver = filters.get('receiver')
        if receiver:
            query += " AND receiver LIKE ?"
            params.append(f'%{receiver}%')
        
        # Apply search query
        search = filters.get('search')
        if search:
            query += " AND (sender LIKE ? OR receiver LIKE ? OR subject LIKE ? OR summary LIKE ?)"
            params.extend([f'%{search}%', f'%{search}%', f'%{search}%', f'%{search}%'])
        
        # Get latest run if no specific date range
        if not date_range and not filters.get('allRuns', False):
            c.execute('SELECT id FROM summary_runs ORDER BY id DESC LIMIT 1')
            latest_run = c.fetchone()
            if latest_run:
                query += " AND run_id = ?"
                params.append(latest_run[0])
        
        query += " ORDER BY email_number"
        
        c.execute(query, params)
        rows = c.fetchall()
        
        email_data = []
        for row in rows:
            email_data.append({
                "number": row[0],
                "from": row[1],
                "to": row[2],
                "subject": row[3],
                "summary": row[4],
                "date": row[5]
            })
        
        conn.close()
        return email_data
        
    except Exception as e:
        print(f"‚ùå Error getting filtered email data: {e}")
        return []

def export_csv(email_data, filename):
    """Export email data as CSV"""
    try:
        # Create CSV in memory
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow(['#', 'Sender', 'Receiver', 'Date', 'Subject', 'Summary'])
        
        # Write data
        for email in email_data:
            writer.writerow([
                email['number'],
                email['from'],
                email['to'],
                email.get('date', ''),
                email['subject'],
                email['summary']
            ])
        
        # Create response
        output.seek(0)
        mem_file = io.BytesIO()
        mem_file.write(output.getvalue().encode('utf-8'))
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='text/csv',
            as_attachment=True,
            download_name=f'{filename}.csv'
        )
        
    except Exception as e:
        print(f"‚ùå CSV export error: {e}")
        raise

def export_json(email_data, filename):
    """Export email data as JSON"""
    try:
        # Create JSON in memory
        mem_file = io.BytesIO()
        mem_file.write(json.dumps(email_data, indent=2).encode('utf-8'))
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='application/json',
            as_attachment=True,
            download_name=f'{filename}.json'
        )
        
    except Exception as e:
        print(f"‚ùå JSON export error: {e}")
        raise

def export_word(email_data, filename):
    """Export email data as Word document"""
    try:
        print("üìÑ Creating Word document...")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Email Summary Report', 0)
        title.alignment = 1
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Emails: {len(email_data)}")
        doc.add_paragraph()
        
        # Add table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['#', 'Sender', 'Receiver', 'Subject', 'Summary']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Data rows
        for email in email_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(email['number'])
            row_cells[1].text = str(email['from'])[:40]
            row_cells[2].text = str(email['to'])[:40]
            row_cells[3].text = str(email['subject'])[:80]
            row_cells[4].text = str(email['summary'])[:400]
        
        # Save to memory
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
        
    except Exception as e:
        print(f"‚ùå Word export error: {e}")
        raise

def export_pdf(email_data, filename):
    """Export email data as PDF"""
    try:
        print("üìÑ Creating PDF document...")
        
        # Create PDF in memory
        mem_file = io.BytesIO()
        doc = SimpleDocTemplate(mem_file, pagesize=letter)
        elements = []
        
        # Add title
        styles = getSampleStyleSheet()
        title = Paragraph("Email Summary Report", styles['Title'])
        elements.append(title)
        
        # Add metadata
        elements.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        elements.append(Paragraph(f"Total Emails: {len(email_data)}", styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # Prepare table data
        table_data = [['#', 'Sender', 'Receiver', 'Subject', 'Summary']]
        
        for email in email_data:
            table_data.append([
                str(email['number']),
                str(email['from'])[:40],
                str(email['to'])[:40],
                str(email['subject'])[:60],
                str(email['summary'])[:100]
            ])
        
        # Create table
        table = Table(table_data, colWidths=[30, 100, 100, 150, 200])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.whitesmoke])
        ]))
        
        elements.append(table)
        
        # Build PDF
        doc.build(elements)
        mem_file.seek(0)
        
        return send_file(
            mem_file,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{filename}.pdf'
        )
        
    except Exception as e:
        print(f"‚ùå PDF export error: {e}")
        raise

@app.route('/api/all-email-data')
@login_required
def get_all_email_data():
    """Get all email data for export"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        c.execute('''
            SELECT email_number, sender, receiver, subject, summary, email_date 
            FROM email_data 
            ORDER BY email_date DESC, email_number
        ''')
        
        rows = c.fetchall()
        conn.close()
        
        email_data = []
        for row in rows:
            email_data.append({
                "number": row[0],
                "from": row[1],
                "to": row[2],
                "subject": row[3],
                "summary": row[4],
                "date": row[5]
            })
        
        return jsonify(email_data)
        
    except Exception as e:
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

# ==================== FLASK ROUTES ====================

@app.route('/')
def root():
    """Redirect root to login or dashboard based on auth"""
    if 'logged_in' in session:
        return redirect('/dashboard')
    return redirect('/login')

@app.route('/dashboard')
@login_required
def dashboard():
    """Main dashboard page"""
    try:
        print(f"üìä Serving dashboard page to user '{session.get('username')}'...")
        return render_template('dashboard.html', username=session.get('username'))
    except Exception as e:
        return f"""
        <html>
            <body style="font-family: Arial, sans-serif; padding: 40px;">
                <h1>‚ùå Dashboard Error</h1>
                <p><strong>Error:</strong> {str(e)}</p>
                <p><a href="/test-html">Test HTML Page</a></p>
                <p><a href="/api">API Endpoints</a></p>
            </body>
        </html>
        """, 500

@app.route('/test-html')
def test_html():
    """Test if HTML rendering works"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Test Page</title>
        <style>
            body { 
                font-family: Arial, sans-serif; 
                padding: 40px;
                background: #f5f7fa;
            }
            .container {
                max-width: 600px;
                margin: 0 auto;
                background: white;
                padding: 30px;
                border-radius: 12px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            }
            .success { 
                color: #10b981; 
                font-size: 24px;
            }
            .btn {
                display: inline-block;
                padding: 12px 24px;
                background: #4f46e5;
                color: white;
                text-decoration: none;
                border-radius: 8px;
                margin: 10px 5px;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1 class="success">‚úÖ HTML is Working!</h1>
            <p>If you can see this page, Flask can render HTML properly.</p>
            <p>Now check the dashboard link below:</p>
            <div>
                <a class="btn" href="/dashboard">Go to Dashboard</a>
                <a class="btn" href="/api">API Endpoints</a>
            </div>
        </div>
    </body>
    </html>
    """

@app.route('/api')
@login_required
def api_home():
    """API home page"""
    return jsonify({
        "status": "Email Summarizer API is running",
        "timestamp": datetime.now().isoformat(),
        "user": session.get('username'),
        "endpoints": {
            "dashboard": "/dashboard",
            "test_page": "/test-html", 
            "health": "/health",
            "stats": "/api/stats",
            "recent_summaries": "/api/recent-summaries",
            "debug": "/api/debug",
            "debug_database": "/api/debug-database",
            "test_json": "/api/test-json",
            "fix_database": "/api/fix-database",
            "force_test_run": "/api/force-test-run",
            "trigger_manual": "/api/trigger-manual (POST)",
            "export_data": "/api/export-data (POST)",
            "all_email_data": "/api/all-email-data"
        }
    })

@app.route('/api/debug')
@login_required
def api_debug():
    """Debug API endpoint"""
    return jsonify({
        "status": "API is working",
        "timestamp": datetime.now().isoformat(),
        "user": session.get('username'),
        "endpoint": "/api/debug"
    })

@app.route('/api/test-json')
@login_required
def test_json():
    """Test JSON response"""
    return jsonify({
        "message": "This is a test JSON response",
        "numbers": [1, 2, 3],
        "timestamp": datetime.now().isoformat(),
        "user": session.get('username')
    })

@app.route('/api/debug-database')
@login_required
def debug_database():
    """Debug database contents"""
    try:
        db_path = get_db_path()
        print(f"üîç Debugging database at: {db_path}")
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Check if tables exist
        c.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = c.fetchall()
        
        # Check summary_runs
        c.execute('SELECT COUNT(*) as run_count FROM summary_runs')
        run_count = c.fetchone()[0]
        
        # Check email_data
        c.execute('SELECT COUNT(*) as email_count FROM email_data')
        email_count = c.fetchone()[0]
        
        # Get latest run details
        c.execute('SELECT * FROM summary_runs ORDER BY id DESC LIMIT 1')
        latest_run = c.fetchone()
        
        # Get some email data samples
        c.execute('SELECT * FROM email_data ORDER BY id DESC LIMIT 5')
        sample_emails = c.fetchall()
        
        conn.close()
        
        return jsonify({
            "database_status": "connected",
            "database_path": db_path,
            "tables_found": [table[0] for table in tables],
            "summary_runs_count": run_count,
            "email_data_count": email_count,
            "latest_run": {
                "id": latest_run[0] if latest_run else None,
                "date": latest_run[1] if latest_run else None,
                "total_emails": latest_run[2] if latest_run else None,
                "processed_emails": latest_run[3] if latest_run else None
            } if latest_run else None,
            "sample_emails": [
                {
                    "id": email[0],
                    "run_id": email[1], 
                    "email_number": email[2],
                    "sender": email[3],
                    "subject": email[5]
                } for email in sample_emails
            ]
        })
        
    except Exception as e:
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/fix-database')
@admin_required
def fix_database():
    """Debug and fix database issues"""
    try:
        print("üîß Fixing database...")
        
        # Reinitialize database
        init_db()
        
        # Add a test run to verify
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Add a test run
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        c.execute('''
            INSERT INTO summary_runs 
            (run_date, total_emails, processed_emails, success_rate, status)
            VALUES (?, ?, ?, ?, ?)
        ''', (current_time, 2, 2, 100.0, 'test'))
        
        run_id = c.lastrowid
        
        # Add test emails
        test_emails = [
            (run_id, 1, "test@example.com", "archives@jubalandstate.so", "Test Email 1", "This is a test summary for email 1.", current_time),
            (run_id, 2, "admin@example.com", "archives@jubalandstate.so", "Test Email 2", "This is a test summary for email 2.", current_time)
        ]
        
        c.executemany('''
            INSERT INTO email_data 
            (run_id, email_number, sender, receiver, subject, summary, email_date)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', test_emails)
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "status": "success",
            "message": "Database fixed and test data added",
            "test_run_id": run_id
        })
        
    except Exception as e:
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/force-test-run')
@admin_required
def force_test_run():
    """Force a test run with sample data"""
    try:
        # Create sample data for testing
        sample_emails = [
            {
                "from": "test@jubalandstate.so",
                "to": "archives@jubalandstate.so", 
                "subject": "Test Email 1",
                "body": "This is a test email body for testing the dashboard."
            },
            {
                "from": "admin@jubalandstate.so",
                "to": "archives@jubalandstate.so",
                "subject": "Test Email 2", 
                "body": "Another test email to verify dashboard functionality."
            }
        ]
        
        sample_summaries = {
            1: "This is a test summary for email 1. It demonstrates how summaries will appear in the dashboard.",
            2: "This is a test summary for email 2. The dashboard should display this data properly."
        }
        
        # Store sample data
        success = store_email_data_for_dashboard(sample_emails, sample_summaries)
        
        return jsonify({
            "status": "success" if success else "error",
            "message": "Test data added to dashboard" if success else "Failed to add test data",
            "emails_added": len(sample_emails)
        })
        
    except Exception as e:
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/stats')
@login_required
def get_stats():
    """API endpoint for dashboard statistics"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        c.execute('''
            SELECT * FROM summary_runs 
            ORDER BY id DESC 
            LIMIT 1
        ''')
        
        latest_run = c.fetchone()
        conn.close()
        
        if latest_run:
            stats = {
                "total_emails_today": latest_run[2] or 0,
                "emails_processed": latest_run[3] or 0,
                "success_rate": round(latest_run[4] or 0, 1),
                "last_run": latest_run[1],
                "next_run": (datetime.now() + timedelta(days=1)).strftime('%Y-%m-%d 09:00:00'),
                "deepseek_usage": "Calculating...",
                "status": "active",
                "user": session.get('username')
            }
        else:
            # Default stats if no runs yet
            stats = {
                "total_emails_today": 0,
                "emails_processed": 0,
                "success_rate": 0,
                "last_run": "Never",
                "next_run": (datetime.now() + timedelta(days=1)).strftime('%Y-%m-%d 09:00:00'),
                "deepseek_usage": "0 tokens",
                "status": "waiting",
                "user": session.get('username')
            }
        
        return jsonify(stats)
    except Exception as e:
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/api/recent-summaries')
@login_required
def get_recent_summaries():
    """API endpoint for recent email summaries - FIXED VERSION"""
    try:
        db_path = get_db_path()
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Get the latest run ID
        c.execute('SELECT id FROM summary_runs ORDER BY id DESC LIMIT 1')
        latest_run = c.fetchone()
        
        if not latest_run:
            print("üì≠ No runs found in database, using fallback data")
            conn.close()
            return jsonify(get_fallback_email_data())
        
        run_id = latest_run[0]
        print(f"üîç Fetching emails for run_id: {run_id}")
        
        # Get ALL email data for the latest run
        c.execute('''
            SELECT email_number, sender, receiver, subject, summary, email_date 
            FROM email_data 
            WHERE run_id = ? 
            ORDER BY email_number
        ''', (run_id,))
        
        email_data = []
        rows = c.fetchall()
        print(f"üìß Found {len(rows)} email records for run_id {run_id}")
        
        for row in rows:
            email_data.append({
                "number": row[0],
                "from": row[1],
                "to": row[2],
                "subject": row[3],
                "summary": row[4],
                "date": row[5]
            })
        
        conn.close()
        
        # If no data found, use fallback
        if not email_data:
            email_data = get_fallback_email_data()
        
        print(f"üìä Returning {len(email_data)} emails for dashboard table")
        return jsonify(email_data)
        
    except Exception as e:
        print(f"‚ùå Error getting recent summaries: {e}")
        print(f"Full traceback: {traceback.format_exc()}")
        # Fallback to mock data if database is not available
        return jsonify(get_fallback_email_data())

def get_fallback_email_data():
    """Provide fallback data if database is not available"""
    return [
        {
            "number": 1,
            "from": "system@jubalandstate.so",
            "to": "archives@jubalandstate.so",
            "subject": "Daily System Report",
            "summary": "Automated system report showing all services are running normally with 99.8% uptime. No critical issues reported.",
            "date": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        },
        {
            "number": 2,
            "from": "secretary@jubalandstate.so", 
            "to": "archives@jubalandstate.so",
            "subject": "Meeting Minutes Approval",
            "summary": "Requesting approval for executive meeting minutes. Key decisions include budget allocation and project timelines.",
            "date": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
    ]

@app.route('/api/trigger-manual', methods=['POST'])
@admin_required
def trigger_manual_run():
    """Manually trigger email summary process"""
    try:
        # Run in background thread to avoid timeout
        def run_background():
            try:
                agent = EmailSummarizerAgent()
                agent.run_complete_summary()
            except Exception as e:
                print(f"‚ùå Error in background run: {e}")
                print(f"Full traceback: {traceback.format_exc()}")
        
        thread = threading.Thread(target=run_background)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            "status": "success", 
            "message": "Email summary process started in background. This may take 10-15 minutes.",
            "timestamp": datetime.now().isoformat(),
            "triggered_by": session.get('username')
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e), "traceback": traceback.format_exc()}), 500

@app.route('/health')
def health():
    return jsonify({"status": "healthy", "timestamp": datetime.now().isoformat()})

# ==================== EMAIL SUMMARIZER CLASS ====================

class EmailSummarizerAgent:
    def __init__(self):
        # Use environment variables for security
        self.deepseek_api_key = os.getenv('DEEPSEEK_API_KEY')
        self.source_email = os.getenv('SOURCE_EMAIL')
        self.source_password = os.getenv('SOURCE_PASSWORD')
        self.imap_server = os.getenv('IMAP_SERVER', 'imap.one.com')
        
        # Validate required environment variables
        if not self.deepseek_api_key:
            raise ValueError("Missing DEEPSEEK_API_KEY environment variable")
        if not self.source_email:
            raise ValueError("Missing SOURCE_EMAIL environment variable")
        if not self.source_password:
            raise ValueError("Missing SOURCE_PASSWORD environment variable")
            
        self.deepseek_api_url = "https://api.deepseek.com/v1/chat/completions"
        self.imap_port = 993
    
    def fetch_emails_last_24h(self):
        try:
            print("üìß Connecting to One.com IMAP server...")
            mail = imaplib.IMAP4_SSL(self.imap_server, self.imap_port)
            mail.login(self.source_email, self.source_password)
            mail.select("inbox")
            
            since_date = (datetime.now() - timedelta(hours=24)).strftime("%d-%b-%Y")
            print(f"üìÖ Fetching emails since: {since_date}")
            
            status, messages = mail.search(None, f'(SINCE "{since_date}")')
            
            if status != 'OK':
                print("üì≠ No emails found")
                mail.close()
                mail.logout()
                return []
                
            email_ids = messages[0].split()
            print(f"‚úÖ Found {len(email_ids)} emails in last 24 hours")
            
            emails_data = []
            
            # Process emails
            for i, email_id in enumerate(email_ids, 1):
                try:
                    email_id_str = email_id.decode('utf-8') if isinstance(email_id, bytes) else str(email_id)
                    
                    status, msg_data = mail.fetch(email_id_str, "(RFC822)")
                    if status != 'OK':
                        continue
                        
                    msg = email.message_from_bytes(msg_data[0][1])
                    
                    subject = self.decode_email_header(msg.get("Subject", ""))
                    from_ = self.decode_email_header(msg.get("From", ""))
                    to_ = self.decode_email_header(msg.get("To", "") or msg.get("Delivered-To", "Unknown"))
                    date = msg.get("Date", "")
                    
                    body = self.extract_email_body(msg)
                    
                    emails_data.append({
                        "subject": subject,
                        "from": from_,
                        "to": to_,
                        "date": date,
                        "body": body[:1000]  # Limit for token management
                    })
                    
                    if i % 10 == 0:
                        print(f"üì• Processed {i}/{len(email_ids)} emails...")
                    
                except Exception as e:
                    print(f"‚ö†Ô∏è Error processing email {email_id}: {e}")
                    continue
            
            mail.close()
            mail.logout()
            
            print(f"‚úÖ Successfully processed {len(emails_data)} emails")
            return emails_data
            
        except Exception as e:
            print(f"‚ùå Error fetching emails: {e}")
            print(f"Full traceback: {traceback.format_exc()}")
            return []
    
    def decode_email_header(self, header):
        if not header:
            return ""
        
        try:
            decoded_parts = decode_header(header)
            decoded_header = ""
            for part, encoding in decoded_parts:
                if isinstance(part, bytes):
                    if encoding:
                        decoded_header += part.decode(encoding, errors='ignore')
                    else:
                        decoded_header += part.decode('utf-8', errors='ignore')
                else:
                    decoded_header += str(part)
            return decoded_header
        except Exception:
            return str(header)
    
    def extract_email_body(self, msg):
        body = ""
        
        try:
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get("Content-Disposition", ""))
                    
                    if content_type == "text/plain" and "attachment" not in content_disposition:
                        try:
                            payload = part.get_payload(decode=True)
                            if payload:
                                body = payload.decode('utf-8', errors='ignore')
                                if body.strip():
                                    break
                        except:
                            continue
            else:
                content_type = msg.get_content_type()
                if content_type == "text/plain":
                    try:
                        payload = msg.get_payload(decode=True)
                        if payload:
                            body = payload.decode('utf-8', errors='ignore')
                    except:
                        pass
        except Exception as e:
            print(f"‚ö†Ô∏è Error extracting email body: {e}")
        
        return body
    
    def summarize_emails_in_batches(self, emails_data):
        """Summarize emails in batches to handle token limits"""
        if not emails_data:
            print("üì≠ No emails to summarize")
            return {}
        
        print(f"üìù Summarizing {len(emails_data)} emails in batches...")
        
        # Process emails in smaller batches
        batch_size = 10  # Reduced from 20 to avoid token limits
        all_summaries = {}
        
        for batch_num in range(0, len(emails_data), batch_size):
            batch_emails = emails_data[batch_num:batch_num + batch_size]
            batch_summaries = self._summarize_batch(batch_emails, batch_num)
            all_summaries.update(batch_summaries)
            
            # Add delay between batches to avoid rate limiting
            if batch_num + batch_size < len(emails_data):
                print(f"‚è≥ Waiting 3 seconds before next batch...")
                time.sleep(3)
        
        return all_summaries
    
    def _summarize_batch(self, batch_emails, start_index):
        """Summarize one batch of emails"""
        if not batch_emails:
            return {}
        
        emails_text = ""
        for i, email in enumerate(batch_emails, 1):
            email_num = start_index + i
            emails_text += f"Email {email_num}:\n"
            emails_text += f"From: {email.get('from', 'Unknown')}\n"
            emails_text += f"To: {email.get('to', 'Unknown')}\n"
            emails_text += f"Subject: {email.get('subject', 'No Subject')}\n"
            emails_text += f"Date: {email.get('date', 'Unknown')}\n"
            emails_text += f"Content: {email.get('body', '')[:150]}...\n\n"
        
        prompt = f"""
        Please provide individual one-paragraph summaries for EACH email. Format your response exactly like this:

        **Email {start_index + 1}:** [One paragraph summary of this email]
        **Email {start_index + 2}:** [One paragraph summary of this email]
        ...and so on for each email.

        Make each summary concise but informative, focusing on the main purpose and key points of each email.
        Keep each summary to 2-3 sentences maximum.

        Emails to summarize ({len(batch_emails)} emails in this batch):
        {emails_text}
        """
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.deepseek_api_key}"
        }
        
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system", 
                    "content": "Provide clear, concise individual one-paragraph summaries for each email. Format each summary starting with **Email X:** followed by the paragraph. Keep summaries brief (2-3 sentences)."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "max_tokens": 2000,  # Reduced for smaller batches
            "temperature": 0.3
        }
        
        try:
            batch_num = (start_index // 10) + 1
            print(f"ü§ñ Summarizing batch {batch_num} ({len(batch_emails)} emails)...")
            
            response = requests.post(self.deepseek_api_url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            
            result = response.json()
            summary_text = result['choices'][0]['message']['content']
            
            # Parse individual summaries
            batch_summaries = self.extract_individual_summaries(summary_text, batch_emails, start_index)
            print(f"‚úÖ Batch {batch_num} summarized successfully")
            
            return batch_summaries
            
        except requests.exceptions.RequestException as e:
            print(f"‚ùå API request failed for batch: {e}")
        except Exception as e:
            print(f"‚ùå Error summarizing batch: {e}")
        
        # Return empty summaries for this batch if failed
        return {start_index + i + 1: "Summary unavailable (API error)" for i in range(len(batch_emails))}
    
    def extract_individual_summaries(self, summary_text, batch_emails, start_index):
        """Extract individual email summaries from batch response"""
        summaries = {}
        
        for i in range(len(batch_emails)):
            email_num = start_index + i + 1
            
            # Try multiple patterns to find the summary
            patterns = [
                r"\*\*Email\s+" + str(email_num) + r":\*\*\s*(.*?)(?=\*\*Email\s+" + str(email_num + 1) + r":\*\*|$)",
                r"Email\s+" + str(email_num) + r":\s*(.*?)(?=Email\s+" + str(email_num + 1) + r":|$)",
            ]
            
            found = False
            for pattern in patterns:
                match = re.search(pattern, summary_text, re.DOTALL | re.IGNORECASE)
                if match:
                    summary = match.group(1).strip()
                    # Clean up the summary
                    summary = re.sub(r'\*\*', '', summary)
                    summary = re.sub(r'\s+', ' ', summary)
                    summary = summary[:400]  # Limit length for table
                    summaries[email_num] = summary
                    found = True
                    break
            
            # Fallback: if no pattern found, try to extract from lines
            if not found:
                lines = summary_text.split('\n')
                for line in lines:
                    line_lower = line.lower()
                    if f"email {email_num}:" in line_lower or f"**email {email_num}:**" in line_lower:
                        summary = line.replace(f"Email {email_num}:", "").replace(f"**Email {email_num}:**", "").replace("**", "").strip()
                        if summary:
                            summaries[email_num] = summary[:400]
                            found = True
                            break
            
            # Final fallback
            if not found:
                summaries[email_num] = "Summary not available"
        
        return summaries
    
    def create_word_document(self, emails_data, all_summaries):
        try:
            print("üìÑ Creating Word document...")
            
            doc = Document()
            title = doc.add_heading('Email Summary Report', 0)
            title.alignment = 1
            
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Period: Last 24 Hours")
            doc.add_paragraph(f"Total Emails Processed: {len(emails_data)}")
            doc.add_paragraph()
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No'
            hdr_cells[1].text = 'Sender'
            hdr_cells[2].text = 'Receiver'
            hdr_cells[3].text = 'Email Subject'
            hdr_cells[4].text = 'Summary in a paragraph'
            
            # Data rows
            for i, email in enumerate(emails_data, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = str(email.get('from', 'Unknown'))[:40]
                row_cells[2].text = str(email.get('to', 'Unknown'))[:40]
                row_cells[3].text = str(email.get('subject', 'No Subject'))[:80]
                
                # Get individual summary for this email
                summary = all_summaries.get(i, "Summary being processed...")
                row_cells[4].text = str(summary)
            
            filename = f"Complete_Email_Summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            print(f"‚úÖ Word document saved: {filename}")
            
            return filename
            
        except Exception as e:
            print(f"‚ùå Error creating Word document: {e}")
            return None
    
    def run_complete_summary(self):
        print(f"\n{'='*60}")
        print(f"üöÄ STARTING COMPLETE EMAIL SUMMARY - {datetime.now()}")
        print(f"{'='*60}")
        
        try:
            # Step 1: Fetch ALL emails from last 24 hours
            emails_data = self.fetch_emails_last_24h()
            
            if not emails_data:
                print("üì≠ No emails to process")
                # Still create an empty run record
                store_email_data_for_dashboard([], {})
                return
                
            print(f"üìß Processing {len(emails_data)} emails...")
            
            # Step 2: Summarize ALL emails in batches
            all_summaries = self.summarize_emails_in_batches(emails_data)
            
            print(f"üìù Generated {len(all_summaries)} summaries out of {len(emails_data)} emails")
            
            # Step 3: Store the processed emails and summaries for the dashboard
            storage_success = store_email_data_for_dashboard(emails_data, all_summaries)
            
            if storage_success:
                print("‚úÖ Email data successfully stored for dashboard")
            else:
                print("‚ùå Failed to store email data for dashboard")
            
            # Step 4: VERIFY DATA STORAGE
            print(f"\n{'='*60}")
            print("üîç VERIFYING DATA STORAGE FOR DASHBOARD...")
            print(f"{'='*60}")
            verify_data_storage()
            
            # Step 5: Create Word document (optional)
            try:
                self.create_word_document(emails_data, all_summaries)
            except Exception as e:
                print(f"‚ö†Ô∏è Word document creation skipped: {e}")
            
            print(f"\n‚úÖ COMPLETE summary process finished at {datetime.now()}")
            print(f"üìä Processed {len(emails_data)} emails total")
            print(f"üìã Generated {len(all_summaries)} summaries")
            print(f"üíæ Data sent to dashboard successfully")
                
        except Exception as e:
            print(f"‚ùå Critical error in complete summary: {e}")
            print(f"Full traceback: {traceback.format_exc()}")

# ==================== DATABASE FUNCTIONS ====================

def store_email_data_for_dashboard(emails_data, all_summaries):
    """Store processed email data for dashboard display - FIXED VERSION"""
    try:
        db_path = get_db_path()
        print(f"üíæ Storing {len(emails_data)} emails in database at: {db_path}")
        
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Create new run entry
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        total_emails = len(emails_data)
        processed_emails = len(all_summaries)
        success_rate = (processed_emails / total_emails * 100) if total_emails > 0 else 0
        
        c.execute('''
            INSERT INTO summary_runs 
            (run_date, total_emails, processed_emails, success_rate, status)
            VALUES (?, ?, ?, ?, ?)
        ''', (
            current_time,
            total_emails,
            processed_emails,
            success_rate,
            'completed'
        ))
        
        run_id = c.lastrowid
        print(f"üìä Created new run_id: {run_id}")
        
        # Insert email data
        inserted_count = 0
        
        for i, email in enumerate(emails_data, 1):
            summary = all_summaries.get(i, "Summary not available")
            
            try:
                c.execute('''
                    INSERT INTO email_data 
                    (run_id, email_number, sender, receiver, subject, summary, email_date)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
                    run_id,
                    i,
                    str(email.get('from', 'Unknown'))[:100],
                    str(email.get('to', 'Unknown'))[:100],
                    str(email.get('subject', 'No Subject'))[:200],
                    str(summary)[:500],
                    email.get('date', current_time)
                ))
                inserted_count += 1
                
            except Exception as e:
                print(f"‚ö†Ô∏è Failed to store email {i}: {e}")
                continue
        
        conn.commit()
        conn.close()
        
        print(f"‚úÖ Database storage complete:")
        print(f"   ‚úÖ Run ID: {run_id}")
        print(f"   ‚úÖ Emails stored: {inserted_count}/{total_emails}")
        print(f"   ‚úÖ Success rate: {success_rate:.1f}%")
        
        return inserted_count > 0
        
    except Exception as e:
        print(f"‚ùå CRITICAL ERROR storing email data: {e}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False

def verify_data_storage():
    """Verify that data was properly stored in database"""
    try:
        db_path = get_db_path()
        print(f"üîç Verifying database at: {db_path}")
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Check latest run
        c.execute('''
            SELECT id, run_date, total_emails, processed_emails 
            FROM summary_runs 
            ORDER BY id DESC LIMIT 1
        ''')
        latest_run = c.fetchone()
        
        if not latest_run:
            print("‚ùå VERIFICATION FAILED: No runs found in database")
            conn.close()
            return False
            
        run_id, run_date, total_emails, processed_emails = latest_run
        print(f"üìã Latest run: ID={run_id}, Date={run_date}, Total Emails={total_emails}, Processed={processed_emails}")
        
        # Check email data for this run
        c.execute('SELECT COUNT(*) FROM email_data WHERE run_id = ?', (run_id,))
        stored_emails = c.fetchone()[0]
        
        print(f"üìã Stored emails for run {run_id}: {stored_emails}")
        
        # Get a sample of stored data
        c.execute('''
            SELECT email_number, sender, receiver, subject 
            FROM email_data 
            WHERE run_id = ? 
            ORDER BY email_number 
            LIMIT 3
        ''', (run_id,))
        samples = c.fetchall()
        
        if samples:
            print("üìã Sample stored emails:")
            for sample in samples:
                print(f"   - #{sample[0]}: From '{sample[1]}' to '{sample[2]}' - '{sample[3]}'")
        else:
            print("üì≠ No email samples found")
        
        conn.close()
        
        success = stored_emails > 0
        if success:
            print(f"‚úÖ VERIFICATION PASSED: {stored_emails} emails stored in database")
        else:
            print(f"‚ùå VERIFICATION FAILED: No emails stored for run {run_id}")
            
        return success
        
    except Exception as e:
        print(f"‚ùå VERIFICATION ERROR: {e}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False

def scheduled_summary():
    """Function to be called by Render Cron Job"""
    try:
        print(f"üïí Running scheduled summary at {datetime.now()}")
        agent = EmailSummarizerAgent()
        agent.run_complete_summary()
        return True
    except Exception as e:
        print(f"‚ùå Scheduled summary failed: {e}")
        print(f"Full traceback: {traceback.format_exc()}")
        return False

# For direct script execution (cron job)
if __name__ == "__main__":
    # Check if running as cron job (you can set CRON_MODE environment variable)
    if os.getenv('CRON_MODE'):
        print("üïí Running in cron mode...")
        scheduled_summary()
    else:
        # Web service mode
        print("üåê Starting web server...")
        app.run(host='0.0.0.0', port=5000, debug=False)
